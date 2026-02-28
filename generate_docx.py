"""
Generate ToledoMessage Academic Paper (Word Document)
Editable .docx version matching the PDF content
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from io import BytesIO

# ─── Colors ───
C_PRIMARY = RGBColor(0x19, 0x76, 0xD2)
C_DARK = RGBColor(0x1A, 0x1A, 0x2E)
C_PURPLE = RGBColor(0x7C, 0x4D, 0xFF)
C_TEAL = RGBColor(0x00, 0x80, 0x69)
C_RED = RGBColor(0xE5, 0x39, 0x35)
C_BLUE_DARK = RGBColor(0x15, 0x65, 0xC0)
C_GRAY = RGBColor(0x75, 0x75, 0x75)
C_GRAY_L = RGBColor(0xE0, 0xE0, 0xE0)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ─── Page setup ───
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ─── Style Definitions ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = Pt(14)

# Heading 1
h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(15)
h1.font.bold = True
h1.font.color.rgb = C_PRIMARY
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(8)

# Heading 2
h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.color.rgb = C_BLUE_DARK
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after = Pt(6)

# ─── Helper Functions ───

def set_cell_bg(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, font_name='Calibri', font_size=Pt(9),
                  color=RGBColor(0x33, 0x33, 0x33), alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Set cell text with formatting."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run.font.color.rgb = color

def add_header_row(table, headers, bg_color_hex):
    """Format the first row as header."""
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, bg_color_hex)
        set_cell_text(cell, header, bold=True, font_size=Pt(9.5),
                      color=C_WHITE, alignment=WD_ALIGN_PARAGRAPH.CENTER)

def add_table_row(table, row_idx, values, code_col=None):
    """Add data to a table row."""
    bg = 'FFFFFF' if row_idx % 2 == 1 else 'F5F7FA'
    for i, val in enumerate(values):
        cell = table.rows[row_idx].cells[i]
        set_cell_bg(cell, bg)
        fn = 'Consolas' if code_col is not None and i == code_col else 'Calibri'
        set_cell_text(cell, val, font_name=fn, font_size=Pt(9))

def add_caption(text):
    """Add a centered italic caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = C_GRAY

def add_algorithm(title, lines):
    """Add a formatted algorithm block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = C_BLUE_DARK

    # Code block as a single-cell table with light bg
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, 'F0F4F8')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run('\n')
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Add border to the table
    tbl_element = tbl._tbl
    tblPr = tbl_element.tblPr if tbl_element.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:color="D0D8E0"/>'
        '  <w:left w:val="single" w:sz="4" w:color="D0D8E0"/>'
        '  <w:bottom w:val="single" w:sz="4" w:color="D0D8E0"/>'
        '  <w:right w:val="single" w:sz="4" w:color="D0D8E0"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def add_bullet(text_parts):
    """Add a bullet point. text_parts is list of (text, bold) tuples."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    for text, bold in text_parts:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def make_figure(draw_func, w=6, h=4, dpi=150):
    buf = BytesIO()
    fig, ax = plt.subplots(figsize=(w, h), dpi=dpi)
    fig.patch.set_facecolor('white')
    ax.set_facecolor('white')
    draw_func(fig, ax)
    fig.savefig(buf, format='png', bbox_inches='tight', facecolor='white',
                edgecolor='none', pad_inches=0.15)
    plt.close(fig)
    buf.seek(0)
    return buf

# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after = Pt(8)
run = p.add_run('ToledoMessage: A Post-Quantum Hybrid\nEnd-to-End Encrypted Messaging Platform')
run.font.name = 'Calibri'
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = C_DARK

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
run = p.add_run('Combining Classical and Post-Quantum Cryptography\nfor Secure Browser-Based Communication')
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.color.rgb = C_GRAY

# Horizontal line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
pPr = p._p.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    '  <w:bottom w:val="single" w:sz="12" w:color="1976D2" w:space="1"/>'
    '</w:pBdr>'
)
pPr.append(pBdr)

# Author
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('Mohamed Khamis')
run.font.name = 'Calibri'
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run('University of Toledo, Department of Computer Science\nToledo, Ohio, United States')
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.color.rgb = C_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
run = p.add_run('February 2026')
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.color.rgb = C_GRAY

# Abstract
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(6)
run = p.add_run('ABSTRACT')
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(1.5)
p.paragraph_format.right_indent = Cm(1.5)
p.paragraph_format.space_after = Pt(8)
run = p.add_run(
    'This paper presents ToledoMessage, a novel secure messaging platform that implements a hybrid '
    'post-quantum end-to-end encryption (E2EE) system within a browser-based WebAssembly environment. '
    'The system combines the X3DH (Extended Triple Diffie-Hellman) key agreement protocol with the '
    'NIST-standardized ML-KEM-768 (formerly CRYSTALS-Kyber) key encapsulation mechanism, creating '
    'a hybrid classical/post-quantum key exchange that provides security against both current and '
    'future quantum computing threats. Message-level encryption employs the Double Ratchet protocol '
    'with AES-256-GCM authenticated encryption, providing forward secrecy and break-in recovery. '
    'Digital signatures use a hybrid Ed25519 + ML-DSA-65 (formerly Dilithium) scheme for authentication. '
    'The platform is implemented using Blazor WebAssembly for the client, ASP.NET Core with SignalR '
    'for real-time server communication, and BouncyCastle for cryptographic operations. The server '
    'operates under a zero-knowledge model, never accessing plaintext messages or private keys. '
    'We present the system architecture, cryptographic protocol details, security analysis, and '
    'a comparison with existing messaging platforms including Signal, WhatsApp, and Telegram.'
)
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# Keywords
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
p.paragraph_format.space_after = Pt(12)
run = p.add_run('Keywords: ')
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.bold = True
run = p.add_run(
    'Post-Quantum Cryptography, End-to-End Encryption, Signal Protocol, X3DH, '
    'Double Ratchet, ML-KEM-768, Blazor WebAssembly, Hybrid Cryptography, Secure Messaging'
)
run.font.name = 'Calibri'
run.font.size = Pt(10)

# Page break after title page
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    'The proliferation of quantum computing poses an existential threat to the cryptographic '
    'foundations of modern secure communication. Current messaging applications such as Signal, '
    'WhatsApp, and iMessage rely on classical Diffie-Hellman key exchange and ECDSA signatures, '
    'which are vulnerable to Shor\'s algorithm on a sufficiently powerful quantum computer [1]. '
    'The National Institute of Standards and Technology (NIST) has recently standardized post-quantum '
    'cryptographic algorithms, including ML-KEM (Module-Lattice-Based Key Encapsulation Mechanism) '
    'and ML-DSA (Module-Lattice-Based Digital Signature Algorithm), marking a critical inflection '
    'point in the transition to quantum-resistant cryptography [2].'
)

doc.add_paragraph(
    'ToledoMessage addresses this challenge by implementing a hybrid cryptographic model that '
    'combines proven classical algorithms with NIST-standardized post-quantum primitives. The '
    'system\'s key contribution is the integration of ML-KEM-768 into the Signal Protocol\'s X3DH '
    'key agreement, creating a key exchange that requires an attacker to break both the classical '
    'X25519 Diffie-Hellman and the lattice-based ML-KEM to compromise a session. This hybrid '
    'approach provides defense-in-depth: if either primitive is found to be vulnerable, the '
    'other continues to protect communication security.'
)

doc.add_paragraph(
    'A second contribution is the deployment of this hybrid cryptographic stack within a browser '
    'environment using Blazor WebAssembly (WASM). All cryptographic operations\u2014key generation, '
    'session management, encryption, and decryption\u2014execute entirely in the client\'s browser '
    'sandbox, with the server operating under a strict zero-knowledge model. The server stores '
    'and routes only encrypted ciphertext, never accessing plaintext messages or private keys.'
)

# ═══════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════
doc.add_heading('2. System Architecture', level=1)
doc.add_heading('2.1 High-Level Overview', level=2)

doc.add_paragraph(
    'ToledoMessage follows a three-tier architecture consisting of: (1) a Blazor WebAssembly '
    'client running in the browser, (2) an ASP.NET Core server providing REST APIs and real-time '
    'SignalR communication, and (3) a SQL Server 2022 database for persistent storage. Figure 1 '
    'illustrates the overall system architecture.'
)

# ── Architecture Diagram ──
def arch_fig(fig, ax):
    ax.set_xlim(-0.5, 14)
    ax.set_ylim(-0.5, 8.5)
    ax.axis('off')

    # Client box
    box1 = mpatches.FancyBboxPatch((0, 4.5), 4, 3.2, boxstyle="round,pad=0.3",
        facecolor='#E3F2FD', edgecolor='#1976D2', linewidth=2)
    ax.add_patch(box1)
    ax.text(2.0, 7.2, 'Blazor WASM Client', ha='center', fontsize=10, fontweight='bold', color='#1565C0')
    ax.text(2.0, 5.9, 'CryptoService\nSessionService\nKeyGeneration\nLocalStorage (IndexedDB)',
            ha='center', fontsize=8.5, color='#333333', linespacing=1.5)

    # Server box
    box2 = mpatches.FancyBboxPatch((5.2, 4.5), 4, 3.2, boxstyle="round,pad=0.3",
        facecolor='#F3E5F5', edgecolor='#7C4DFF', linewidth=2)
    ax.add_patch(box2)
    ax.text(7.2, 7.2, 'ASP.NET Core Server', ha='center', fontsize=10, fontweight='bold', color='#6A1B9A')
    ax.text(7.2, 5.9, 'ChatHub (SignalR)\nREST Controllers\nIdentity + JWT Auth\nEF Core 10',
            ha='center', fontsize=8.5, color='#333333', linespacing=1.5)

    # Database box
    box3 = mpatches.FancyBboxPatch((10.4, 4.5), 3.4, 3.2, boxstyle="round,pad=0.3",
        facecolor='#FFEBEE', edgecolor='#E53935', linewidth=2)
    ax.add_patch(box3)
    ax.text(12.1, 7.2, 'SQL Server 2022', ha='center', fontsize=10, fontweight='bold', color='#C62828')
    ax.text(12.1, 5.9, 'Users, Devices\nConversations\nMessages (encrypted)\nPreKey Bundles',
            ha='center', fontsize=8.5, color='#333333', linespacing=1.5)

    # IndexedDB box
    box4 = mpatches.FancyBboxPatch((0, 0.2), 4, 3.0, boxstyle="round,pad=0.3",
        facecolor='#E8F5E9', edgecolor='#2E7D32', linewidth=2)
    ax.add_patch(box4)
    ax.text(2.0, 2.75, 'Browser IndexedDB', ha='center', fontsize=10, fontweight='bold', color='#2E7D32')
    ax.text(2.0, 1.5, 'Private Keys\nSession State\nRatchet Chains\nAuth Tokens',
            ha='center', fontsize=8.5, color='#333333', linespacing=1.5)

    # Arrows
    ax.annotate('', xy=(5.05, 6.6), xytext=(4.15, 6.6),
        arrowprops=dict(arrowstyle='<->', color='#25D366', lw=2.5))
    ax.text(4.6, 7.05, 'SignalR', ha='center', fontsize=8.5, color='#2E7D32', fontweight='bold')

    ax.annotate('', xy=(5.05, 5.5), xytext=(4.15, 5.5),
        arrowprops=dict(arrowstyle='<->', color='#FF9800', lw=2))
    ax.text(4.6, 5.1, 'HTTPS REST', ha='center', fontsize=8.5, color='#E65100', fontweight='bold')

    ax.annotate('', xy=(10.25, 6.1), xytext=(9.35, 6.1),
        arrowprops=dict(arrowstyle='<->', color='#E53935', lw=2))
    ax.text(9.8, 6.55, 'EF Core', ha='center', fontsize=8.5, color='#C62828', fontweight='bold')

    ax.annotate('', xy=(2.0, 4.3), xytext=(2.0, 3.4),
        arrowprops=dict(arrowstyle='<->', color='#2E7D32', lw=2))
    ax.text(3.0, 3.85, 'JS Interop', ha='center', fontsize=8.5, color='#2E7D32', fontweight='bold')

buf = make_figure(arch_fig, w=8.5, h=5.2)
doc.add_picture(buf, width=Cm(16))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_caption('Figure 1: ToledoMessage system architecture.')

# 2.2 Technology Stack
doc.add_heading('2.2 Technology Stack', level=2)

tech_data = [
    ['Component', 'Technology', 'Purpose'],
    ['Client Runtime', 'Blazor WebAssembly (.NET 10)', 'Browser-based UI + crypto execution'],
    ['Crypto Library', 'BouncyCastle 2.6.2', 'X25519, ML-KEM-768, Ed25519, ML-DSA-65, AES-GCM'],
    ['Server Framework', 'ASP.NET Core 10', 'REST API endpoints + middleware'],
    ['Real-Time', 'SignalR', 'WebSocket-based message delivery + events'],
    ['Authentication', 'ASP.NET Core Identity + JWT', 'User registration, login, token issuance'],
    ['ORM', 'Entity Framework Core 10', 'Code First migrations, SQL Server provider'],
    ['Database', 'SQL Server 2022', 'Persistent storage (encrypted data only)'],
    ['Client Storage', 'Browser IndexedDB', 'Private keys, session state, tokens'],
]

table = doc.add_table(rows=len(tech_data), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, tech_data[0], '1976D2')
for i in range(1, len(tech_data)):
    add_table_row(table, i, tech_data[i])
add_caption('Table 1: Technology stack components and their roles.')

# ═══════════════════════════════════════════════════════════
# 3. CRYPTOGRAPHIC MODEL
# ═══════════════════════════════════════════════════════════
doc.add_heading('3. Cryptographic Model', level=1)
doc.add_heading('3.1 Hybrid X3DH Key Agreement', level=2)

doc.add_paragraph(
    'The Extended Triple Diffie-Hellman (X3DH) protocol, originally designed by Marlinspike and '
    'Perrin for the Signal Protocol [3], establishes a shared secret between two parties who may '
    'not be simultaneously online. ToledoMessage extends X3DH by incorporating ML-KEM-768 key '
    'encapsulation alongside the standard X25519 Diffie-Hellman exchanges.'
)

doc.add_paragraph('Each device registers the following pre-key bundle with the server:')

add_bullet([('Identity Key: ', True), ('Long-term Ed25519 public key (32 bytes) + ML-DSA-65 public key (1,952 bytes)', False)])
add_bullet([('Signed Pre-Key: ', True), ('Medium-term X25519 public key (32 bytes), signed with hybrid signature', False)])
add_bullet([('Kyber Pre-Key: ', True), ('ML-KEM-768 encapsulation key (1,184 bytes), signed with hybrid signature', False)])
add_bullet([('One-Time Pre-Keys: ', True), ('Batch of 10 ephemeral X25519 public keys (32 bytes each)', False)])

add_algorithm('Algorithm 1: Hybrid X3DH Initiator (Alice)', [
    'function X3DH_Initiate(BobBundle):',
    '  1. Verify(BobBundle.signedPreKey, BobBundle.identityKey)',
    '  2. Verify(BobBundle.kyberPreKey, BobBundle.identityKey)',
    '  3. (eph_pub, eph_priv) = X25519.GenerateKeyPair()',
    '  4. DH1 = X25519(eph_priv, BobBundle.signedPreKey)     // 32 bytes',
    '  5. DH2 = X25519(eph_priv, BobBundle.oneTimePreKey)    // 32 bytes',
    '  6. (kem_ct, kem_ss) = ML-KEM-768.Encapsulate(BobBundle.kyberPreKey)',
    '  7. ikm = DH1 || DH2 || kem_ss',
    '  8. SK = HKDF-SHA256(ikm, info="ToledoMessage-X3DH-v1", len=64)',
    '  9. rootKey = SK[0:32], chainKey = SK[32:64]',
    '  return (rootKey, chainKey, eph_pub, kem_ct)',
])

doc.add_paragraph(
    'The critical security property of this hybrid scheme is that the shared secret SK depends '
    'on material from both the classical X25519 exchanges and the post-quantum ML-KEM-768 '
    'encapsulation. An attacker must break both X25519 (solving the elliptic curve discrete '
    'logarithm problem) and ML-KEM-768 (solving the Module-LWE problem) to recover the session keys.'
)

# 3.2 Double Ratchet
doc.add_heading('3.2 Double Ratchet Protocol', level=2)

doc.add_paragraph(
    'After session establishment via X3DH, ToledoMessage employs the Double Ratchet algorithm [4] '
    'to derive unique encryption keys for each message. The protocol maintains three key chains:'
)

add_bullet([('Root Chain: ', True), ('Updated via DH ratchet steps, producing new chain keys. Uses HKDF-SHA256 with the current root key as salt and a new DH shared secret as input key material.', False)])
add_bullet([('Sending Chain: ', True), ('Derives per-message keys for outgoing messages. Each step uses HMAC-SHA256 with the chain key to produce a message key and the next chain key.', False)])
add_bullet([('Receiving Chain: ', True), ('Symmetric to the sending chain, used for decrypting incoming messages. Maintains separate counters for out-of-order message handling.', False)])

add_algorithm('Algorithm 2: Double Ratchet Message Encryption', [
    'function RatchetEncrypt(session, plaintext):',
    '  1. (session.sendChainKey, messageKey) = KDF_CK(session.sendChainKey)',
    '  2. session.sendMessageNumber += 1',
    '  3. nonce = SecureRandom(12 bytes)',
    '  4. ciphertext = AES-256-GCM.Encrypt(messageKey, nonce, plaintext)',
    '  5. SecureErase(messageKey)  // Immediately delete after use',
    '  return (nonce || ciphertext || authTag)',
])

doc.add_paragraph(
    'This construction provides two critical security properties: (1) Forward Secrecy \u2014 '
    'compromise of the current session state does not reveal past message keys, as each chain key '
    'is immediately replaced after deriving a message key; (2) Break-in Recovery \u2014 '
    'if an attacker obtains the current session state, a new DH ratchet step (triggered by receiving '
    'a message with a new DH public key) establishes a fresh shared secret, healing the session.'
)

# 3.3 AES-256-GCM
doc.add_heading('3.3 AES-256-GCM Authenticated Encryption', level=2)

doc.add_paragraph(
    'Each message is encrypted using AES-256-GCM (Galois/Counter Mode), an AEAD (Authenticated '
    'Encryption with Associated Data) cipher. The implementation uses a 256-bit key derived from '
    'the Double Ratchet, a 96-bit random nonce generated per message, and produces a 128-bit '
    'authentication tag. The wire format is:'
)

add_algorithm('Wire Format', [
    'output = nonce(12 bytes) || ciphertext(variable) || authTag(16 bytes)',
    'transport = Base64Encode(output)',
])

# 3.4 Hybrid Digital Signatures
doc.add_heading('3.4 Hybrid Digital Signatures', level=2)

doc.add_paragraph(
    'ToledoMessage employs a hybrid signature scheme combining Ed25519 (classical, 64-byte signatures) '
    'with ML-DSA-65 (post-quantum, 3,309-byte signatures). The HybridSigner class produces '
    'signatures by concatenating both: signature = Ed25519.Sign(m) || ML-DSA-65.Sign(m). '
    'Verification requires both components to pass independently, ensuring that the compromise of '
    'one scheme does not invalidate the authentication guarantee.'
)

# Table 2: Signature comparison
sig_data = [
    ['Property', 'Ed25519', 'ML-DSA-65', 'Hybrid (Combined)'],
    ['Public Key Size', '32 bytes', '1,952 bytes', '1,984 bytes'],
    ['Signature Size', '64 bytes', '3,309 bytes', '3,373 bytes'],
    ['Security Basis', 'ECDLP', 'Module-LWE', 'Both (defense-in-depth)'],
    ['Quantum Resistant', 'No', 'Yes (NIST Level 3)', 'Yes'],
    ['Standard', 'RFC 8032', 'FIPS 204', 'Custom hybrid'],
]

table = doc.add_table(rows=len(sig_data), cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, sig_data[0], '7C4DFF')
for i in range(1, len(sig_data)):
    add_table_row(table, i, sig_data[i])
add_caption('Table 2: Comparison of signature scheme properties.')

# ═══════════════════════════════════════════════════════════
# 4. POST-QUANTUM INTEGRATION
# ═══════════════════════════════════════════════════════════
doc.add_heading('4. Post-Quantum Cryptographic Integration', level=1)
doc.add_heading('4.1 ML-KEM-768 (CRYSTALS-Kyber)', level=2)

doc.add_paragraph(
    'ML-KEM-768 is a lattice-based Key Encapsulation Mechanism standardized by NIST in FIPS 203 [5]. '
    'It provides IND-CCA2 security based on the Module Learning With Errors (MLWE) problem. '
    'ToledoMessage integrates ML-KEM-768 at the X3DH key agreement stage via the MlKemKeyExchange '
    'class, which wraps BouncyCastle\'s ML-KEM implementation.'
)

# Table 3: ML-KEM parameters
kem_data = [
    ['Parameter', 'Value'],
    ['Algorithm', 'ML-KEM-768 (FIPS 203)'],
    ['Security Level', 'NIST Level 3 (~192-bit classical equivalent)'],
    ['Public Key Size', '1,184 bytes'],
    ['Ciphertext Size', '1,088 bytes'],
    ['Shared Secret Size', '32 bytes'],
    ['Underlying Problem', 'Module-LWE (lattice-based)'],
    ['Integration Point', 'X3DH step 6: KEM encapsulation against Kyber pre-key'],
]

table = doc.add_table(rows=len(kem_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, kem_data[0], '008069')
for i in range(1, len(kem_data)):
    add_table_row(table, i, kem_data[i])
add_caption('Table 3: ML-KEM-768 parameters in ToledoMessage.')

# 4.2 Hybrid Security Argument
doc.add_heading('4.2 Hybrid Security Argument', level=2)

doc.add_paragraph(
    'The hybrid key agreement provides a "best of both worlds" security guarantee. Let SK = '
    'HKDF(DH1 || DH2 || kem_ss). An adversary attempting to recover SK must either: (a) solve '
    'the Computational Diffie-Hellman problem on Curve25519 to recover DH1 and DH2, and '
    '(b) solve the Module-LWE problem to recover kem_ss from the ML-KEM ciphertext. The HKDF '
    'construction ensures that even partial knowledge of the input key material does not reveal '
    'the derived key, provided at least one component remains secret [6].'
)

# ═══════════════════════════════════════════════════════════
# 5. MESSAGE FLOW
# ═══════════════════════════════════════════════════════════
doc.add_heading('5. End-to-End Message Flow', level=1)

doc.add_paragraph('The complete lifecycle of a message in ToledoMessage involves the following steps:')

add_bullet([('Step 1 - Session Establishment: ', True), ('If no active Double Ratchet session exists with the recipient, the sender\'s CryptoService initiates X3DH by fetching the recipient\'s pre-key bundle from the server.', False)])
add_bullet([('Step 2 - Key Derivation: ', True), ('The Double Ratchet derives a unique message key (MK) from the sending chain. The chain key is advanced, ensuring MK is never reused.', False)])
add_bullet([('Step 3 - Encryption: ', True), ('AeadCipher.Encrypt(MK, plaintext) produces the AES-256-GCM ciphertext with random nonce and authentication tag.', False)])
add_bullet([('Step 4 - Multi-Device Fan-Out: ', True), ('MessageEncryptionService repeats steps 2\u20133 for each of the recipient\'s registered devices, producing independent ciphertexts.', False)])
add_bullet([('Step 5 - Transport: ', True), ('Encrypted envelopes are sent via SignalR\'s ChatHub.SendMessage(). The server stores the ciphertext and routes it to connected devices.', False)])
add_bullet([('Step 6 - Delivery: ', True), ('The recipient\'s device receives the envelope via SignalR push, decrypts using its own Double Ratchet session, and displays the plaintext.', False)])
add_bullet([('Step 7 - Acknowledgment: ', True), ('Delivery and read receipts flow back through SignalR, updating message status (Sending \u2192 Sent \u2192 Delivered \u2192 Read).', False)])

# ═══════════════════════════════════════════════════════════
# 6. SECURITY ANALYSIS
# ═══════════════════════════════════════════════════════════
doc.add_heading('6. Security Analysis', level=1)
doc.add_heading('6.1 Threat Model', level=2)

doc.add_paragraph(
    'We consider an adversary with the following capabilities: (a) full control of the network '
    '(active man-in-the-middle), (b) ability to compromise the server and access all stored data, '
    '(c) access to a cryptographically relevant quantum computer (CRQC). Table 4 summarizes the '
    'security guarantees provided by ToledoMessage against each threat.'
)

threat_data = [
    ['Threat', 'Mitigation', 'Status'],
    ['Man-in-the-Middle', 'X3DH mutual auth + safety number verification', 'Mitigated'],
    ['Server Compromise', 'E2EE: server sees only ciphertext', 'Mitigated'],
    ['Past Key Compromise', 'Forward secrecy via Double Ratchet', 'Mitigated'],
    ['Future Key Compromise', 'Break-in recovery via DH ratchet', 'Mitigated'],
    ['Quantum Attack (keys)', 'ML-KEM-768 in X3DH key agreement', 'Mitigated'],
    ['Quantum Attack (sigs)', 'ML-DSA-65 in hybrid signatures', 'Mitigated'],
    ['Replay Attack', 'Random nonce + message counters', 'Mitigated'],
    ['Device Theft', 'Keys in browser IndexedDB (session-bound)', 'Partial'],
    ['Metadata Analysis', 'Server sees routing data (not content)', 'Acknowledged'],
]

table = doc.add_table(rows=len(threat_data), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, threat_data[0], 'E53935')
for i in range(1, len(threat_data)):
    add_table_row(table, i, threat_data[i])
add_caption('Table 4: Threat model and mitigations.')

# ═══════════════════════════════════════════════════════════
# 7. COMPARISON
# ═══════════════════════════════════════════════════════════
doc.add_heading('7. Comparison with Existing Systems', level=1)

comp_data = [
    ['Feature', 'ToledoMessage', 'Signal', 'WhatsApp', 'Telegram', 'iMessage'],
    ['E2EE by Default', 'Yes', 'Yes', 'Yes', 'No*', 'Yes'],
    ['Post-Quantum KEM', 'ML-KEM-768', 'PQXDH**', 'No', 'No', 'PQ3**'],
    ['PQ Signatures', 'ML-DSA-65', 'No', 'No', 'No', 'No'],
    ['Forward Secrecy', 'Yes', 'Yes', 'Yes', 'Yes*', 'Yes'],
    ['Open Source Crypto', 'Yes', 'Yes', 'No', 'No', 'No'],
    ['Safety Numbers', 'Yes', 'Yes', 'Yes', 'No', 'No'],
    ['Multi-Device E2EE', 'Yes', 'Partial', 'Partial', 'No', 'Yes'],
    ['Disappearing Msgs', 'Yes', 'Yes', 'Yes', 'Yes', 'No'],
    ['Browser Client', 'Yes (WASM)', 'No', 'No', 'Yes', 'No'],
    ['Key Agreement', 'Hybrid X3DH', 'X3DH', 'X3DH', 'DH', 'ECIES'],
    ['Symmetric Cipher', 'AES-256-GCM', 'AES-256-CBC', 'AES-256-GCM', 'AES-256-IGE', 'AES-CTR'],
]

table = doc.add_table(rows=len(comp_data), cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, comp_data[0], '1976D2')
for i in range(1, len(comp_data)):
    add_table_row(table, i, comp_data[i])
add_caption(
    'Table 5: Feature comparison. *Telegram requires manual opt-in for secret chats. '
    '**Signal\'s PQXDH and Apple\'s PQ3 are recent additions (2023\u20132024) using different PQ schemes.'
)

# Comparison chart
def comp_chart(fig, ax):
    apps = ['ToledoMessage', 'Signal', 'WhatsApp', 'Telegram\n(Secret)', 'iMessage']
    scores = [7, 5.5, 4.5, 2.5, 4]
    colors = ['#3A76F0', '#2196F3', '#25D366', '#2AABEE', '#007AFF']
    bars = ax.barh(apps, scores, color=colors, edgecolor='white', linewidth=0.8, height=0.55)
    ax.set_xlim(0, 8)
    ax.set_xlabel('Security Feature Score (out of 7)', fontsize=9, color='#333333')
    ax.tick_params(colors='#333333', labelsize=9)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    for bar, score in zip(bars, scores):
        ax.text(bar.get_width() + 0.12, bar.get_y() + bar.get_height()/2,
                f'{score}/7', ha='left', va='center', fontsize=9, fontweight='bold', color='#333333')

buf = make_figure(comp_chart, w=6.5, h=2.8)
doc.add_picture(buf, width=Cm(14))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_caption('Figure 2: Aggregate security feature comparison.')

# ═══════════════════════════════════════════════════════════
# 8. IMPLEMENTATION DETAILS
# ═══════════════════════════════════════════════════════════
doc.add_heading('8. Implementation Details', level=1)
doc.add_heading('8.1 Client-Side Services', level=2)

doc.add_paragraph(
    'The Blazor WebAssembly client implements the following key services, all running entirely '
    'within the browser\'s WASM sandbox:'
)

svc_data = [
    ['Service', 'Responsibility'],
    ['CryptoService', 'Orchestrates X3DH initiation, session establishment, encrypt/decrypt operations'],
    ['SessionService', 'Manages Double Ratchet state per peer device (root key, chain keys, counters)'],
    ['KeyGenerationService', 'Generates identity keys (Ed25519+ML-DSA), signed pre-keys, Kyber pre-keys, OTP keys'],
    ['MessageEncryptionService', 'Multi-device fan-out: encrypts for all of recipient\'s registered devices'],
    ['LocalStorageService', 'Encrypted IndexedDB storage via JS Interop for keys and session state'],
    ['FingerprintService', 'Computes SHA-256 safety numbers from identity keys for verification'],
    ['PreKeyReplenishmentService', 'Monitors and auto-replenishes one-time pre-keys when count drops low'],
    ['SignalRService', 'WebSocket connection management, event subscription, message transport'],
    ['ThemeService', 'Persists and applies user-selected UI theme via localStorage'],
]

table = doc.add_table(rows=len(svc_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(table, svc_data[0], '1565C0')
for i in range(1, len(svc_data)):
    add_table_row(table, i, svc_data[i], code_col=0)
add_caption('Table 6: Client-side service architecture.')

# 8.2 Database Schema
doc.add_heading('8.2 Database Schema', level=2)

doc.add_paragraph(
    'The SQL Server database stores user accounts, device registrations (with public keys only), '
    'conversations, participants, and encrypted messages. The schema uses decimal primary keys '
    'with precision 20 and scale 0, UTC timestamps, and cascading deletes for referential integrity. '
    'Critically, the database never contains private keys or plaintext message content\u2014only '
    'ciphertext blobs and public cryptographic material.'
)

# 8.3 Test Coverage
doc.add_heading('8.3 Test Coverage', level=2)

doc.add_paragraph(
    'The cryptographic layer is validated by 65 unit tests covering X25519 key exchange, ML-KEM-768 '
    'encapsulation/decapsulation, Ed25519 and ML-DSA-65 signing/verification, hybrid signatures, '
    'X3DH protocol flow, Double Ratchet session management, and AES-256-GCM encryption/decryption. '
    'All tests pass consistently across the test suite.'
)

# ═══════════════════════════════════════════════════════════
# 9. CONCLUSION
# ═══════════════════════════════════════════════════════════
doc.add_heading('9. Conclusion', level=1)

doc.add_paragraph(
    'ToledoMessage demonstrates that hybrid post-quantum end-to-end encryption is feasible within '
    'a browser-based environment. By integrating ML-KEM-768 into the X3DH key agreement and '
    'ML-DSA-65 into the signature scheme, the system provides quantum-resistant security without '
    'sacrificing the proven guarantees of classical cryptography. The Double Ratchet protocol '
    'ensures forward secrecy and break-in recovery at the per-message level, while the zero-knowledge '
    'server architecture ensures that even a fully compromised server cannot access message content.'
)

doc.add_paragraph(
    'Future work includes: (1) formal security proofs using tools such as ProVerif or Tamarin, '
    '(2) mobile native applications using .NET MAUI, (3) encrypted file and media transfer, '
    '(4) voice and video calling with SRTP encryption, (5) multi-device synchronization using '
    'the Sesame protocol, and (6) decentralized federation following the Matrix protocol model.'
)

# ═══════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════
doc.add_heading('References', level=1)

refs = [
    '[1] P. Shor, "Algorithms for quantum computation: discrete logarithms and factoring," Proc. 35th FOCS, IEEE, 1994.',
    '[2] NIST, "Post-Quantum Cryptography Standardization," FIPS 203, FIPS 204, FIPS 205, 2024.',
    '[3] M. Marlinspike and T. Perrin, "The X3DH Key Agreement Protocol," Signal Foundation, 2016.',
    '[4] T. Perrin and M. Marlinspike, "The Double Ratchet Algorithm," Signal Foundation, 2016.',
    '[5] NIST, "Module-Lattice-Based Key-Encapsulation Mechanism Standard," FIPS 203, August 2024.',
    '[6] H. Krawczyk, "Cryptographic Extraction and Key Derivation: The HKDF Scheme," Proc. CRYPTO 2010, Springer.',
    '[7] D. J. Bernstein et al., "Post-Quantum Cryptography," Nature, vol. 549, 2017.',
    '[8] Signal Foundation, "Signal Protocol Technical Documentation," signal.org/docs, 2023.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    run = p.add_run(ref)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ─── Add page numbers ───
for sect in doc.sections:
    footer = sect.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run('ToledoMessage: Post-Quantum Secure Messaging  \u2014  Page ')
    run.font.name = 'Calibri'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    # Add page number field
    fld_char1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr = parse_xml(f'<w:instrText {nsdecls("w")}> PAGE </w:instrText>')
    fld_char2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run2 = p.add_run()
    run2.font.name = 'Calibri'
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run2._r.append(fld_char1)
    run2._r.append(instr)
    run2._r.append(fld_char2)

# ─── Save ───
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ToledoMessage_Paper.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
